import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from fpdf import FPDF

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Çok dilli destekli TTF font yolu
FONT_PATH = os.path.join(os.path.dirname(__file__), '..', 'fonts', 'DejaVuSans.ttf')
FONT_NAME = "DejaVuSans"


def convert_pdf_to_text(content: bytes, filename: str) -> str:
    """PDF → TXT (PyMuPDF ile, Unicode destekli)"""
    text = ""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            text += page.get_text()
    except Exception as e:
        # Hata durumunda boş metin veya bir hata mesajı döndür
        print(f"PDF metin çıkarma hatası: {e}")
        text = ""

    base_name = os.path.splitext(filename)[0]
    output_path = os.path.join(OUTPUT_DIR, f"{base_name}.txt")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

    return output_path


def convert_pdf_to_word(pdf_bytes: bytes, filename: str) -> str:
    """PDF → DOCX (Unicode destekli)"""
    word_doc = Document()
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in doc:
            text = page.get_text("text", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
            if text.strip():
                word_doc.add_paragraph(text)
    except Exception as e:
        print(f"PDF'ten Word'e dönüştürme hatası: {e}")
        # Hata durumunda en azından bir mesaj ekle
        word_doc.add_paragraph(f"[Dönüştürme hatası: {e}]")


    base_name = os.path.splitext(filename)[0]
    output_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")
    word_doc.save(output_path)
    return output_path


def convert_text_to_pdf(content: bytes, filename: str) -> str:
    """TXT → PDF (Unicode destekli - TTF fontlu)"""
    text = content.decode("utf-8")
    base_name = os.path.splitext(filename)[0]
    output_path = os.path.join(OUTPUT_DIR, f"{base_name}.pdf")

    pdf = FPDF()
    pdf.add_page()

    # DejaVuSans fontunu ekle
    pdf.add_font(FONT_NAME, "", FONT_PATH, uni=True)
    pdf.set_font(FONT_NAME, size=12)

    pdf.write(8, text)

    pdf.output(output_path)
    return output_path
