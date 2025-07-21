import os
import tempfile
import pandas as pd
from docx import Document
import io

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def convert_excel_to_word(content: bytes, filename: str) -> str:
    tmp_path = None
    try:
        # Excel içeriğini geçici dosyaya yaz
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # Excel dosyasını oku (Unicode uyumlu)
        df = pd.read_excel(tmp_path, engine='openpyxl')

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

    # Word belgesine yaz (her satır bir paragraf)
    doc = Document()
    # DataFrame'in başlıklarını ekle
    header = ', '.join(str(col) for col in df.columns)
    doc.add_paragraph(header)
    
    for _, row in df.iterrows():
        # Hücreleri Unicode uyumlu şekilde birleştir
        paragraph = ', '.join(str(cell) if not pd.isna(cell) else '' for cell in row.values)
        doc.add_paragraph(paragraph)

    # Kaydet
    base_name = os.path.splitext(filename)[0]
    output_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")
    doc.save(output_path)

    return output_path


def convert_word_to_excel(content: bytes, filename: str) -> str:
    tmp_path = None
    try:
        # Word içeriğini geçici dosyaya yaz
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # Word içeriğini oku (Unicode destekli)
        doc = Document(tmp_path)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


    # Paragrafları satır olarak al
    data = []
    for para in doc.paragraphs:
        if para.text.strip():
            # Virgül ile ayırarak hücrelere böl
            data.append([cell.strip() for cell in para.text.split(',')])

    if not data:
        return "" # veya uygun bir hata yönetimi

    df = pd.DataFrame(data[1:], columns=data[0])

    # Excel olarak kaydet
    base_name = os.path.splitext(filename)[0]
    output_path = os.path.join(OUTPUT_DIR, f"{base_name}.xlsx")
    df.to_excel(output_path, index=False, engine='openpyxl')

    return output_path
